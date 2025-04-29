# import streamlit as st 
# import os
# import base64
# import zipfile
# import json
# from openai import OpenAI
# from docx import Document
# from docx.shared import Inches
# from dotenv import load_dotenv
# from io import BytesIO
# import httpx
# import time

# # Set Streamlit page configuration
# st.set_page_config(page_title="Chakshu News Summarizer", layout="wide")

# # Load environment variables
# load_dotenv()
# #api_key = os.getenv("OPENAI_API_KEY")
# api_key = os.getenv("OPENAI_API_KEY", st.secrets.get("OPENAI_API_KEY"))

# # Initialize OpenAI client
# client = OpenAI(
#     api_key=api_key,
#     http_client=httpx.Client(verify=False)
# )

# def image_to_base64(image_bytes):
#     """Convert image bytes to Base64 encoding."""
#     return base64.b64encode(image_bytes).decode('utf-8')

# def process_images(image_files):
#     """Process images and return a Word document."""
#     doc = Document()
#     doc.add_heading("News Reports", level=1)
#     total_images = len(image_files)
#     progress_bar = st.progress(0)
    
#     for idx, image_file in enumerate(image_files, start=1):
#         try:
#             st.write(f"Processing image {idx}/{total_images}")
#             image_bytes = image_file.read()
#             image_base64 = image_to_base64(image_bytes)
            
#             response = client.chat.completions.create(
#                 model="gpt-4o",
#                 response_format={"type": "json_object"},
#                 messages=[
#                     {
#                         "role": "user",
#                         "content": [
#                             {"type": "text", "text": "Extract the text from this image and identify the headline (largest font size line). Return JSON with 'headline' and 'text'."},
#                             {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_base64}"}}
#                         ],
#                     }
#                 ],
#                 max_tokens=1000,
#             )
            
#             response_content = response.choices[0].message.content
#             json_data = json.loads(response_content)
#             headline = json_data.get("headline", "No headline identified")
#             extracted_text = json_data.get("text", "")
            
#             paragraphs = [p.strip() for p in extracted_text.split("\n") if p.strip()]
#             article_text = "\n\n".join(paragraphs)
#             original_word_count = len(article_text.split())
#             target_word_count = int(original_word_count * 0.5)
            
#             if original_word_count > 0:
#                 summary_response = client.chat.completions.create(
#                     model="gpt-4o",
#                     messages=[{
#                         "role": "user",
#                         "content": f"""
#                         You are an AI summarization assistant. Process the following article:
#                         Headline: {headline}
                        
#                         {article_text}
                        
#                         **The language of the summarized text must be the same as the language of extracted text.**
                        
#                         Please read and understand the entire article below and then produce a contextual summary that preserves all the key points and the core meaning.
#                         The summary should have less than {target_word_count} words (i.e. less than 50% of the original article's word count).
                        
#                         The AI shall write 5 key points from the news in complete meaningful sentence format.
#                         The key points have to be written in full sentences, and there have to be 5 complete meaningful sentences.
                        
#                         **The summary should be written in two paragraphs.** Each paragraph should be between 150 and 200 words.
#                         """
#                     }],
#                     max_tokens=1000
#                 )
#                 summarized_text = summary_response.choices[0].message.content.strip()
#             else:
#                 summarized_text = "No text to summarize."
            
#             # Append results to document
#             doc.add_heading(f"Article {idx}", level=1)
#             doc.add_paragraph("Image Source: Uploaded File")
            
#             image_stream = BytesIO(image_bytes)
#             doc.add_picture(image_stream, width=Inches(4))
            
#             doc.add_heading("Headline", level=2)
#             doc.add_paragraph(headline)
#             doc.add_heading("Summarized Text", level=2)
#             doc.add_paragraph(summarized_text)
#             doc.add_page_break()
            
#             progress_bar.progress(int((idx / total_images) * 100))
#             time.sleep(0.5)
        
#         except Exception as e:
#             st.error(f"Error processing image: {e}")
    
#     output_stream = BytesIO()
#     doc.save(output_stream)
#     output_stream.seek(0)
#     return output_stream

# # Predefined background image
# def get_base64_image(image_path):
#     """Convert an image to Base64 encoding."""
#     with open(image_path, "rb") as img_file:
#         return base64.b64encode(img_file.read()).decode()

# background_image_path = "Slide1.JPG"  # Change this to your local background image path
# background_base64 = get_base64_image(background_image_path)

# # Styling for improved readability
# st.markdown(f"""
#     <style>
#     .stApp {{
#         background-image: url("data:image/jpeg;base64,{background_base64}");
#         background-size: cover;
#         background-position: center;
#         background-attachment: fixed;
#     }}

#     /* White transparent box for the title */
#     .title-box {{
#         background-color: rgba(255, 255, 255, 0.7);
#         padding: 10px;
#         border-radius: 4px;
#         text-align: left;
#         color: black;
#         text-shadow: none;
#         width: 50%;
#         margin-left: 0%
#         margin-top: 10px;
#         margin-bottom: 10px;
#         max-height: 100px; /* Ensures it doesn't expand too much */
#     }}

#     /* Adjusting all UI elements for better readability */
#     .stRadio div[role="radiogroup"], .stFileUploader, .stButton {{
#         background-color: rgba(255, 255, 255, 0.7);
#         padding: 5px;
#         border-radius: 5px;
#         width: 50%;
#     }}

#     .stRadio div[role="radiogroup"] label {{
#         background-color: rgba(255, 255, 255, 0.5);
#         padding: 5px;
#         border-radius: 5px;
#     }}
#     </style>
#     """, unsafe_allow_html=True)

# # Navigation - Select Processing Mode
# st.markdown('<div class="title-box"><h1>Chakshu News Summarizer</h1></div>', unsafe_allow_html=True)

# processing_mode = st.radio("Choose an option:", ["Single Image Processing", "Bulk Image Processing (ZIP)"])

# if processing_mode == "Single Image Processing":
#     uploaded_image = st.file_uploader("Upload a JPG or PNG image", type=["jpg", "png"])

#     if uploaded_image:
#         if st.button("Process Image"):
#             doc_file = process_images([uploaded_image])
#             st.download_button("Download Summary", doc_file, file_name="news_summary.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# elif processing_mode == "Bulk Image Processing (ZIP)":
#     uploaded_zip = st.file_uploader("Upload a ZIP file containing images (JPG, PNG only)", type=["zip"])

#     if uploaded_zip:
#         with zipfile.ZipFile(uploaded_zip, "r") as zip_ref:
#             image_files = [
#                 zip_ref.open(file) for file in zip_ref.namelist()
#                 if file.lower().endswith((".png", ".jpg",".jpeg"))
#             ]
            
#             if image_files:
#                 if st.button("Process Images"):
#                     doc_file = process_images(image_files)
#                     st.download_button("Download Summary", doc_file, file_name="news_summary.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
#             else:
#                 st.error("No valid image files found in the ZIP. Only JPG and PNG are allowed.")
####################zzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzz
import os
import zipfile
import streamlit as st
from PIL import Image
from io import BytesIO
import google.generativeai as genai
from docx import Document
from docx.shared import Inches
import json
import re

# ✅ Set your Gemini API key
genai.configure(api_key="AIzaSyBBAinN0d7m_hB7ZLOr0XYfYedXlUVoMBI")  # Replace with your key

# 🔵 Original function for analyzing a newspaper image
def analyze_newspaper_image(image):
    prompt = """
You are an expert in analyzing newspaper clippings. From the newspaper image, extract the following structured information:
1. News Brand Name (e.g., 'Dainik Bhaskar', 'The Times of India')
2. Heading
3. Subheading (if any)
4. Callout Boxes (if any)(visually distinct boxed or highlighted texts. avoid adding any other words or explanation. Please keep the text as it is in original news.)
5. Date (if mentioned)
6. Location (usually before the main news)
7. News Bureau (e.g., 'New Delhi Bureau')
8. Body Text (remaining text that forms the main article body. Please avoid photo captions)
9. A concise summary of the article (within 120 words). Start the summary by repeating the original headline exactly, from next paragraph a brief summary in your own words. The language of summary should be same as the language of the original text.
10. Sentiment (The sentiment of the news in one word 'Positive' or 'Negative')
Return the result in this JSON format:
{
  "news_brand": "",
  "heading": "",
  "subheading": "",
  "callout_boxes": [],
  "date": "",
  "location": "",
  "news_bureau": "",
  "body_text": "",
  "summary": "",
  "sentiment": ""
}
"""
    model = genai.GenerativeModel("gemini-2.5-pro-exp-03-25")
    response = model.generate_content([prompt, image])
    return response.text

# 🔵 Format the extracted JSON into Word Document
def format_json_to_doc(doc, extracted_json_text):
    try:
        json_text_match = re.search(r"\{[\s\S]*\}", extracted_json_text)
        structured_data = json.loads(json_text_match.group())
    except:
        doc.add_paragraph("❌ Failed to parse structured JSON from Gemini response.")
        return

    doc.add_heading("📰 Extracted Newspaper Article", level=1)

    def add_section(label, content, style="Normal"):
        if content:
            doc.add_heading(label, level=2)
            if isinstance(content, list):
                for item in content:
                    doc.add_paragraph(f"• {item}", style=style)
            else:
                doc.add_paragraph(content.strip(), style=style)

    add_section("Heading", structured_data.get("heading"))
    add_section("Sentiment", structured_data.get("sentiment"))
    add_section("Summary", structured_data.get("summary"))

    heading = structured_data.get("heading")
    subheading = structured_data.get("subheading")
    body_text = structured_data.get("body_text")
    if any([heading, subheading, body_text]):
        doc.add_heading("Body Text", level=2)
        if heading:
            doc.add_paragraph(heading.strip())
        if subheading:
            doc.add_paragraph(subheading.strip())
        if body_text:
            doc.add_paragraph(body_text.strip())

    add_section("Callout Boxes", structured_data.get("callout_boxes", []))
    add_section("News Brand", structured_data.get("news_brand"))

# 🔵 Save single image and extracted text into Word
def save_to_word_with_image(image_file, extracted_text, doc=None):
    if doc is None:
        doc = Document()
    try:
        img = Image.open(image_file)
        buffer = BytesIO()
        img.save(buffer, format="JPEG")
        buffer.seek(0)
        doc.add_picture(buffer, width=Inches(5.5))
        doc.add_paragraph("⬆️ Original Newspaper Image", style="Caption")
    except Exception as e:
        doc.add_paragraph(f"❌ Failed to insert image: {e}")

    format_json_to_doc(doc, extracted_text)
    return doc

# 🔵 Streamlit UI
st.set_page_config(page_title="Newspaper Analyzer", layout="centered")
st.title("📰 Newspaper Image Analyzer")

# 👉 Mode Selection
mode = st.radio("Choose Mode", ("Single Image Search", "Bulk Image Search"))

if mode == "Single Image Search":
    uploaded_file = st.file_uploader("Upload Newspaper Image", type=["jpg", "jpeg", "png"])

    if uploaded_file:
        image = Image.open(uploaded_file)
        st.image(image, caption="🖼️ Uploaded Newspaper", use_container_width=True)

        if st.button("🔍 Analyze Image"):
            with st.spinner("Analyzing..."):
                extracted_text = analyze_newspaper_image(image)
                try:
                    json_text_match = re.search(r"\{[\s\S]*\}", extracted_text)
                    structured = json.loads(json_text_match.group())
                except:
                    structured = None

            if structured:
                st.success("✅ Extraction complete")
                st.subheader("📌 Extracted Details")
                st.markdown(f"**Headline:** {structured.get('heading', 'N/A')}")
                st.markdown(f"**Sentiment:** {structured.get('sentiment', 'N/A')}")
                st.markdown(f"**Summary:** {structured.get('summary', 'N/A')}")
                
                st.subheader("Body Text")
                if structured.get("heading"):
                    st.markdown(f"Headline: {structured['heading']}")
                if structured.get("subheading"):
                    st.markdown(f"Subheading: {structured['subheading']}")
                if structured.get("body_text"):
                    st.markdown(structured["body_text"])
                st.markdown(f"News Brand: {structured.get('news_brand', 'N/A')}")
                st.markdown(f"Callout Boxes: {structured.get('callout_boxes', '[]')}")

                word_file = BytesIO()
                save_to_word_with_image(uploaded_file, extracted_text).save(word_file)
                word_file.seek(0)
                st.download_button("📥 Download Word File", word_file, file_name="Extracted_News.docx")
            else:
                st.error("❌ Could not parse structured response.")

elif mode == "Bulk Image Search":
    uploaded_zip = st.file_uploader("Upload ZIP of Newspaper Images", type=["zip"])

    if uploaded_zip:
        if st.button("🔍 Analyze Bulk Images"):
            with st.spinner("Analyzing images..."):
                temp_folder = "temp_images"
                os.makedirs(temp_folder, exist_ok=True)

                # Extract ZIP
                with zipfile.ZipFile(uploaded_zip, "r") as zip_ref:
                    zip_ref.extractall(temp_folder)

                image_files = [os.path.join(temp_folder, f) for f in os.listdir(temp_folder)
                               if f.lower().endswith((".jpg", ".jpeg", ".png"))]

                if not image_files:
                    st.error("❌ No valid image files found in ZIP.")
                else:
                    doc = Document()
                    for img_path in image_files:
                        try:
                            img = Image.open(img_path)
                            extracted_text = analyze_newspaper_image(img)
                            save_to_word_with_image(img_path, extracted_text, doc)
                        except Exception as e:
                            doc.add_paragraph(f"⚠️ Failed to process {img_path}: {e}")

                    # Finalize Word file
                    output = BytesIO()
                    doc.save(output)
                    output.seek(0)

                    st.success(f"✅ Processed {len(image_files)} images.")
                    st.download_button("📥 Download Combined Word File", output, file_name="Bulk_Extracted_News.docx")

                # Cleanup temp folder
                try:
                    for f in os.listdir(temp_folder):
                        os.remove(os.path.join(temp_folder, f))
                    os.rmdir(temp_folder)
                except Exception as cleanup_error:
                    print(f"Cleanup error: {cleanup_error}")
